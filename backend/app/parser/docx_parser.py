"""DOCX 解析器（python-docx）：按 Heading 样式构建章节树，输出与 PDF 一致的 Document IR。

DOCX 无"物理页码"概念，source_pages 以章节序号代替（来源追溯粒度到章节）。
"""
import math

from docx import Document as DocxDocument

from app.core.errors import PPTError
from app.core.logging import get_logger

logger = get_logger(__name__)


def _is_heading1(par) -> bool:
    name = (par.style.name or "").lower()
    return name.startswith("heading 1") or name.startswith("标题 1")


def parse_docx(path: str, filename: str = "") -> dict:
    try:
        doc = DocxDocument(path)
    except Exception as e:
        logger.error("DOCX 打开失败: %s", e)
        raise PPTError("E2002", str(e)) from e

    paragraphs = [p for p in doc.paragraphs if p.text.strip()]
    char_count = sum(len(p.text) for p in paragraphs)
    table_count = len(doc.tables)
    image_count = len(doc.inline_shapes)

    if char_count < 50 and image_count > 0:
        # 只有图片没有正文 → 与扫描件同等处理
        raise PPTError("E2003", f"DOCX 正文字符数仅 {char_count}")
    if char_count < 50:
        raise PPTError("E2003", f"DOCX 无有效正文({char_count} 字符)")

    # 按 Heading 1 切章节
    chapters: list[dict] = []
    current = {"title": "", "lines": []}
    for p in paragraphs:
        if _is_heading1(p):
            if current["lines"] or current["title"]:
                chapters.append(current)
            current = {"title": p.text.strip()[:40], "lines": []}
        else:
            current["lines"].append(p.text.strip())
    if current["lines"] or current["title"]:
        chapters.append(current)

    # 表格内容并入所在文档末章（简化处理，保证数字可被回查）
    table_texts = []
    for t in doc.tables:
        rows = ["\t".join(c.text.strip() for c in row.cells) for row in t.rows]
        table_texts.append("\n".join(rows))
    if table_texts and chapters:
        chapters[-1]["lines"].append("\n".join(table_texts))

    result_chapters = []
    for i, ch in enumerate(chapters):
        if not ch["title"]:
            ch["title"] = f"第{i + 1}部分"
        result_chapters.append({
            "title": ch["title"],
            "content": "\n".join(ch["lines"]),
            "source_pages": [i + 1],  # DOCX 无页码，用章节序号代替
        })

    # 章节过少时按段落均分（兜底）
    if len(result_chapters) < 2:
        all_lines = [p.text.strip() for p in paragraphs]
        n = max(2, min(6, math.ceil(len(all_lines) / 40)))
        size = math.ceil(len(all_lines) / n)
        result_chapters = [{
            "title": f"第{idx + 1}部分",
            "content": "\n".join(all_lines[i:i + size]),
            "source_pages": [idx + 1],
        } for idx, i in enumerate(range(0, len(all_lines), size))]
        logger.info("DOCX 未识别到 Heading 结构，按段落均分为 %d 个伪章节(兜底)", len(result_chapters))

    # 文档标题去掉扩展名（封面标题不应出现 .docx/.pdf）
    import os as _os
    clean_name = _os.path.splitext(filename)[0] if filename else ""
    title = clean_name or (result_chapters[0]["title"] if result_chapters else "未命名文档")
    ir = {
        "title": title, "page_count": len(result_chapters), "char_count": char_count,
        "image_count": image_count, "table_count": table_count, "is_scanned": False,
        "pages": [{"page": i + 1, "text": c["content"]} for i, c in enumerate(result_chapters)],
        "chapters": result_chapters,
    }
    logger.info("DOCX 解析完成：%d 字 / %d 章节 / %d 表格 / %d 图片",
                char_count, len(result_chapters), table_count, image_count)
    return ir
